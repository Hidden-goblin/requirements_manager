# -*- coding: utf-8 -*-
from docx import Document
import xlsxwriter
from PyQt5.QtWidgets import QFileDialog, QMessageBox
from tinydb import Query


def export_repository(parent = None):
    export_file_name = QFileDialog.getSaveFileName(parent, 'Save repository to...', "", "Word Document (*.docx)")[0]
    #TODO: check file name
    document = Document()
    # styles = document.styles
    # print(repr(WD_BUILTIN_STYLE))
    # print(repr(styles[WD_BUILTIN_STYLE.TABLE_MEDIUM_GRID_3]))
    document.add_heading("Repository", 0)
    channels = set([f['Channel'] for f in parent.repository.get_functionalities()])
    for channel in channels:
        document.add_heading(channel, 1)
        for functionality in parent.repository.search_functionalities(query = Query().Channel == channel):
            document.add_heading(functionality['Name'], 2)
            document.add_heading("Description", 3)
            document.add_paragraph(functionality['Description'])
            document.add_heading("Use cases", 3)
            for usecase in parent.repository.search_use_cases(query = Query()['Functionality link'] == functionality['ID']):
                table = document.add_table(rows = 8, cols=2)
                table.style = 'MediumGrid3-Accent3'
                table.cell(0, 0).text = "Name"
                table.cell(0, 1).text = usecase['Name']
                table.cell(1, 0).text = "ID"
                table.cell(1, 1).text = usecase['ID']
                table.cell(2, 0).text = "Short Description"
                table.cell(2, 1).text = usecase['Short Description']
                table.cell(3, 0).text = "Pre Requisite"
                table.cell(3, 1).text = usecase['Pre Requisite']
                table.cell(4, 0).text = "Steps"
                table.cell(4, 1).text = usecase['Steps']
                table.cell(5, 0).text = "Post Conditions"
                table.cell(5, 1).text = usecase['Post Condition']
                table.cell(6, 0).text = "Exceptional Behaviour"
                table.cell(6, 1).text = usecase['Exceptional Behaviour']
                table.cell(7, 0).text = "Since Release"
                table.cell(7, 1).text = usecase['Release']
                document.add_paragraph('')

        document.add_page_break()
    document.save(export_file_name)
    end_export("The whole repository is exported to '{}'".format(export_file_name))


def export_usecase(parent = None):
    export_file_name = QFileDialog.getSaveFileName(parent, 'Save repository to...', "", "Excel spreadsheet (*.xlsx)")[0]
    workbook = xlsxwriter.Workbook(export_file_name)
    text_wrap = workbook.add_format()
    text_wrap.set_text_wrap()
    text_wrap.set_align("top")
    align_top = workbook.add_format()
    align_top.set_align("top")
    worksheet = workbook.add_worksheet()
    worksheet.hide_gridlines(0)
    channels = set([f['Channel'] for f in parent.repository.get_functionalities()])
    #spreadsheet headers
    worksheet.write(3, 1, "Folder")
    worksheet.write(3, 2, "Requirement Name")
    worksheet.write(3, 3, "Requirement Code")
    worksheet.write(3, 4, "Test Title")
    worksheet.write(3, 5, "Est")
    worksheet.write(3, 6, "Status")
    worksheet.write(3, 7, "Test Step")
    worksheet.write(3, 8, "Expected Result")
    worksheet.write(3, 9, "Priority")
    worksheet.write(3, 10, "Type")
    worksheet.write(3, 11, "Flow ID")
    worksheet.write(3, 12, "Pre Requisite")
    #spreadsheet data
    row_count = 4
    for channel in channels:
        worksheet.write(row_count, 1, channel, align_top)
        for functionality in parent.repository.search_functionalities(query = Query().Channel == channel):
            worksheet.write(row_count, 2, functionality['Name'], align_top)
            for usecase in parent.repository.search_use_cases(query = Query()['Functionality link'] == functionality['ID']):
                worksheet.write(row_count, 4, usecase['Name'], align_top)
                worksheet.write(row_count, 10, usecase['Test Type'], align_top)
                worksheet.write_rich_string(row_count, 12, usecase['Pre Requisite'], text_wrap)
                for step in usecase['Steps'].split("\n"):
                    worksheet.write(row_count, 7, step, text_wrap)
                    row_count += 1
                worksheet.write(row_count - 1, 8, usecase['Post Condition'], text_wrap)
    workbook.close()
    end_export("The test plan output is at '{}'".format(export_file_name))


def end_export(export_message = "Unknonw"):
    msg = QMessageBox()
    msg.setIcon(QMessageBox.Information)
    msg.setText(export_message)
    msg.setWindowTitle("Export succeed")
    msg.setStandardButtons(QMessageBox.Ok)
    msg.exec_()
